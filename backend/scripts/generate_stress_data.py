import os
import random
import string
from reportlab.pdfgen import canvas
from docx import Document
from reportlab.lib.pagesizes import letter

VIDEO_GAMES_TEXT = [
    "Elden Ring is an action role-playing game developed by FromSoftware.",
    "The Legend of Zelda: Breath of the Wild is an open-world action-adventure game.",
    "Minecraft is a sandbox game developed by Mojang Studios.",
    "Grand Theft Auto V is an action-adventure game developed by Rockstar North.",
    "The Witcher 3: Wild Hunt is purely based on the novels by Andrzej Sapkowski.",
    "God of War follows Kratos and his son Atreus in the world of Norse mythology.",
    "Hollow Knight is a Metroidvania game developed by Team Cherry.",
    "Stardew Valley is a simulation role-playing video game developed by Eric Barone.",
    "Celeste is a platform game developed by Maddy Makes Games.",
    "Hades is a roguelike action dungeon crawler developed by Supergiant Games."
]

def generate_random_text(lines=10):
    return "\n".join([random.choice(VIDEO_GAMES_TEXT) + " " + "".join(random.choices(string.ascii_letters, k=20)) for _ in range(lines)])

def create_pdfs(count=10, output_dir="data"):
    for i in range(count):
        filename = os.path.join(output_dir, f"stress_test_doc_{i}.pdf")
        c = canvas.Canvas(filename, pagesize=letter)
        c.drawString(100, 750, f"Stress Test PDF Document {i}")
        text = c.beginText(40, 680)
        for line in generate_random_text(20).split('\n'):
            text.textLine(line)
        c.drawText(text)
        c.save()
        print(f"Created {filename}")

def create_docxs(count=5, output_dir="data"):
    for i in range(count):
        filename = os.path.join(output_dir, f"stress_test_doc_{i}.docx")
        doc = Document()
        doc.add_heading(f'Stress Test DOCX Document {i}', 0)
        doc.add_paragraph(generate_random_text(15))
        doc.save(filename)
        print(f"Created {filename}")

def create_txts(count=5, output_dir="data"):
    for i in range(count):
        filename = os.path.join(output_dir, f"stress_test_doc_{i}.txt")
        with open(filename, "w") as f:
            f.write(f"Stress Test TXT Document {i}\n\n")
            f.write(generate_random_text(30))
        print(f"Created {filename}")

def main():
    if not os.path.exists("data"):
        os.makedirs("data")
    
    print("Generating Stress Test Data...")
    create_pdfs(10)
    create_docxs(5)
    create_txts(5)
    print("Done generating data.")

if __name__ == "__main__":
    main()
